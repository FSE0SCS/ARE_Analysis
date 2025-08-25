import streamlit as st
import pandas as pd
import plotly.express as px
import io
from fpdf import FPDF
from docx import Document
import xlsxwriter

st.set_page_config(layout="wide")

# --- Funciones de Exportación ---
def to_excel(df):
    """Convierte un DataFrame a un archivo Excel en memoria."""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Resultados')
    processed_data = output.getvalue()
    return processed_data

def to_word(df):
    """Crea un documento de Word con los resultados."""
    document = Document()
    document.add_heading('Informe de Análisis de Datos', level=1)
    document.add_paragraph('Este informe presenta un análisis detallado de los datos a partir de las columnas seleccionadas.')
    
    table = document.add_table(df.shape[0] + 1, df.shape[1])
    for j in range(df.shape[1]):
        table.cell(0, j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])
    
    output = io.BytesIO()
    document.save(output)
    processed_data = output.getvalue()
    return processed_data

def to_pdf(df, figure=None):
    """Crea un PDF con la tabla y el gráfico."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="Informe de Análisis de Datos", ln=True, align='C')
    pdf.ln(10)
    
    if figure:
        img_buffer = io.BytesIO()
        figure.write_image(img_buffer, format="png")
        img_buffer.seek(0)
        pdf.image(img_buffer, x=10, y=pdf.get_y(), w=180)
        pdf.ln(100)
    
    pdf.set_font("Arial", 'B', 10)
    col_widths = [40] * len(df.columns)
    
    for i, col in enumerate(df.columns):
        pdf.cell(col_widths[i], 10, str(col), 1, 0, 'C')
    pdf.ln()
    
    pdf.set_font("Arial", '', 8)
    for row in df.itertuples(index=False):
        for cell in row:
            pdf.cell(40, 10, str(cell), 1, 0, 'C')
        pdf.ln()
    
    output = io.BytesIO()
    pdf.output(output)
    processed_data = output.getvalue()
    return processed_data

# --- Interfaz de Streamlit ---
st.title('Analizador de Datos de Excel 📊')
st.markdown("Sube un archivo de Excel con tus datos para comenzar el análisis.")

# --- Inicializar estado de sesión ---
if 'file_uploaded' not in st.session_state:
    st.session_state.file_uploaded = False
    st.session_state.analyze_button_clicked = False
    st.session_state.df = None
    st.session_state.selected_columns = []

# Botón para reiniciar toda la aplicación
if st.button("Hacer otro análisis"):
    st.session_state.clear()
    st.experimental_rerun()

# --- Uploader de archivos ---
uploaded_file = st.file_uploader("Sube tu archivo de Excel", type=['xlsx'], key="file_uploader")

if uploaded_file:
    st.session_state.file_uploaded = True
    st.session_state.analyze_button_clicked = False
    try:
        df = pd.read_excel(uploaded_file, sheet_name='Hoja1')
        df = df.dropna(axis=1, how='all')
        st.session_state.df = df
    except Exception as e:
        st.error(f"Ocurrió un error al leer el archivo Excel. Asegúrate de que tenga una hoja llamada 'Hoja1'. Error: {e}")
        st.session_state.df = None

# --- Contenido principal si se ha subido un archivo ---
if st.session_state.file_uploaded and st.session_state.df is not None:
    st.subheader("Paso 1: Selecciona las columnas para analizar")
    column_names = st.session_state.df.columns.tolist()
    
    selected_columns = st.multiselect(
        "Elige las columnas para tu análisis:",
        options=column_names,
        key="column_selector"
    )

    st.subheader("Paso 2: Haz clic para ejecutar el análisis")
    if st.button("Analizar Datos"):
        st.session_state.analyze_button_clicked = True
        st.session_state.selected_columns = selected_columns

# --- Lógica de análisis y resultados (se ejecuta solo al presionar el botón) ---
if st.session_state.analyze_button_clicked:
    selected_columns_to_analyze = st.session_state.selected_columns
    df_to_analyze = st.session_state.df
    
    if not selected_columns_to_analyze:
        st.warning("Por favor, selecciona al menos una columna antes de analizar.")
    else:
        economic_column_options = [col for col in selected_columns_to_analyze if any(c in str(col).lower() for c in ['euro', '€', 'coste', 'importe', 'valor', 'ingreso', 'precio'])]
        
        if not economic_column_options:
            st.error("No se pudo identificar una columna de valores económicos entre las seleccionadas. Elige una que contenga términos como 'Euro', '€', 'Valor', 'Importe', etc.")
        else:
            economic_column = economic_column_options[0]
            if len(economic_column_options) > 1:
                st.subheader("Selección de columna económica")
                economic_column = st.selectbox(
                    "Se encontraron múltiples columnas económicas. Selecciona la que deseas usar:",
                    options=economic_column_options,
                    key="economic_column_selector"
                )
            
            group_by_columns = [col for col in selected_columns_to_analyze if col != economic_column]
            
            if not group_by_columns:
                st.header("Análisis de la Columna Económica")
                st.subheader(f"Suma total de {economic_column}")
                total_sum = df_to_analyze[economic_column].sum()
                st.metric(label="Suma Total", value=f"€{total_sum:,.2f}")
            else:
                st.header("Resultados del Análisis")
                try:
                    analysis_df = df_to_analyze.groupby(group_by_columns)[economic_column].sum().reset_index()
                    
                    st.subheader("Tabla de Datos Analizados")
                    st.dataframe(analysis_df, use_container_width=True)

                    st.subheader("Gráfico de Resultados")
                    fig = px.bar(analysis_df, x=group_by_columns[0], y=economic_column,
                                 title=f'Suma de {economic_column} por {group_by_columns[0]}',
                                 color=group_by_columns[0] if len(group_by_columns) > 1 else None)
                    st.plotly_chart(fig, use_container_width=True)

                    st.subheader("Resumen del Análisis")
                    total_sum = analysis_df[economic_column].sum()
                    st.markdown(f"""
                    El análisis ha sumado los valores de la columna **{economic_column}** agrupados por **{', '.join(group_by_columns)}**.
                    El total acumulado es de **€{total_sum:,.2f}**.
                    """)
                    
                    st.subheader("Opciones de Exportación")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button(
                            label="📥 Descargar Excel",
                            data=to_excel(analysis_df),
                            file_name='informe_analisis.xlsx',
                            mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                        )
                    with col2:
                        st.download_button(
                            label="📥 Descargar Word",
                            data=to_word(analysis_df),
                            file_name='informe_analisis.docx',
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                    with col3:
                        st.download_button(
                            label="📥 Descargar PDF",
                            data=to_pdf(analysis_df, fig),
                            file_name='informe_analisis.pdf',
                            mime='application/pdf'
                        )
                except Exception as e:
                    st.error(f"Ocurrió un error al procesar los datos. Por favor, verifica tus selecciones y el formato del archivo. Error: {e}")